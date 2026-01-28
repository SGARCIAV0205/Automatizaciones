from pathlib import Path
from datetime import date
from docx import Document
import streamlit as st

def generate_minutes(
    title: str,
    date: date,
    objectives: str,
    summary: str,
    attendees,          # pandas.DataFrame o None
    action_items: list, # lista de dicts
    output_path: Path,
    output_format: str = "docx"  # "docx" o "md"
):
    output_path.parent.mkdir(exist_ok=True, parents=True)

    if output_format == "docx":
        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"Fecha: {date}")

        if objectives:
            doc.add_heading("Objetivos", level=2)
            doc.add_paragraph(objectives)

        if attendees is not None and len(attendees) > 0:
            doc.add_heading("Participantes", level=2)
            t = doc.add_table(rows=1, cols=min(3, len(attendees.columns)))
            hdr = t.rows[0].cells
            cols = attendees.columns.tolist()[: len(hdr)]
            for i, c in enumerate(cols):
                hdr[i].text = str(c)
            for _, row in attendees.iterrows():
                r = t.add_row().cells
                for i, c in enumerate(cols):
                    r[i].text = str(row[c])

        if summary:
            doc.add_heading("Resumen", level=2)
            doc.add_paragraph(summary)

        if action_items:
            doc.add_heading("Acuerdos y Tareas", level=2)
            for i, it in enumerate(action_items, start=1):
                desc = it.get("descripcion","").strip()
                resp = it.get("responsable","").strip()
                due  = it.get("fecha","").strip()
                p = doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                p.add_run(desc)
                if resp or due:
                    doc.add_paragraph(f"Responsable: {resp} | Compromiso: {due}")

        doc.save(output_path)
        return True

    if output_format == "md":
        lines = [f"# {title}", f"**Fecha:** {date}", ""]
        if objectives:
            lines += ["## Objetivos", objectives, ""]
        if attendees is not None and len(attendees) > 0:
            cols = attendees.columns.tolist()
            header = "| " + " | ".join(cols) + " |"
            sep = "| " + " | ".join(["---"]*len(cols)) + " |"
            lines += ["## Participantes", header, sep]
            for _, r in attendees.iterrows():
                lines.append("| " + " | ".join(str(r[c]) for c in cols) + " |")
            lines.append("")
        if summary:
            lines += ["## Resumen", summary, ""]
        if action_items:
            lines.append("## Acuerdos y Tareas")
            for i, it in enumerate(action_items, start=1):
                desc = it.get("descripcion","").strip()
                resp = it.get("responsable","").strip()
                due  = it.get("fecha","").strip()
                lines.append(f"{i}. {desc}")
                if resp or due:
                    lines.append(f"   - Responsable: {resp} | Compromiso: {due}")
        output_path.write_text("\n".join(lines), encoding="utf-8")
        return True

    return False

def render_ai_enhancement_section():
    """Renderizar sección de mejora con AI para minutas"""
    try:
        from modules.openai_client import openai_client, is_openai_available
        
        if not is_openai_available():
            return None, None
        
        st.subheader("Mejora con AI")
        st.info("Usa ChatGPT para generar automáticamente objetivos, resumen y tareas desde una transcripción.")
        
        # Input de transcripción
        transcription = st.text_area(
            "Transcripción de la reunión",
            height=200,
            placeholder="Pega aquí la transcripción de tu reunión...",
            help="El AI analizará este texto para generar la estructura de la minuta"
        )
        
        # Contexto adicional
        context = st.text_input(
            "Contexto adicional (opcional)",
            placeholder="Ej: Reunión de planificación trimestral del equipo de marketing",
            help="Información adicional que ayude al AI a entender mejor el contexto"
        )
        
        if st.button("Generar con AI", type="primary"):
            if not transcription.strip():
                st.warning("Por favor ingresa una transcripción")
                return None, None
            
            with st.spinner("Analizando transcripción con AI..."):
                enhanced_content, error = openai_client.enhance_minutes(transcription, context)
                
                if error:
                    st.error(f"Error al procesar con AI: {error}")
                    return None, None
                else:
                    st.success("Contenido generado con AI")
                    return enhanced_content, transcription
        
        return None, None
        
    except ImportError:
        return None, None
