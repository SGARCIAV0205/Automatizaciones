from typing import Dict
from docx import Document


def _replace_in_paragraph(paragraph, mapping: Dict[str, str]) -> None:
    # Funciona bien si el placeholder no está partido en runs.
    for run in paragraph.runs:
        text = run.text
        for k, v in mapping.items():
            token = "{{" + k + "}}"
            if token in text:
                text = text.replace(token, v)
        run.text = text


def fill_docx_template(template_path: str, output_path: str, mapping: Dict[str, str]) -> None:
    doc = Document(template_path)

    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)

    doc.save(output_path)
